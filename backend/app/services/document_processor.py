import base64
import logging
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import csv
import io
import subprocess

import anthropic
import docx
import fitz
import openpyxl

from app.core.config import settings

logger = logging.getLogger(__name__)

# --- PDF prompts ---

PAGE_EXTRACTION_PROMPT = (
    "このページの内容を構造化Markdownとして忠実に抽出してください。\n"
    "ルール:\n"
    "- 見出し構造（#, ##, ###）を保持する\n"
    "- 表はMarkdownテーブルに変換し、全ての数値と対応関係を正確に再現する\n"
    "- グラフ・チャートは、読み取れるデータポイントを表やリストとして記述する\n"
    "- 図や画像の内容を言語化して記述する\n"
    "- テキストは原文のまま忠実に抽出する\n"
    "- 余計な説明や前置きは不要。抽出結果のみを返す"
)

DOCUMENT_EXTRACTION_PROMPT = (
    "このドキュメントの内容を構造化Markdownとして抽出してください。\n"
    "ルール:\n"
    "- 見出し構造（#, ##, ###）を保持する\n"
    "- 表はMarkdownテーブルに変換する\n"
    "- 図や画像の内容を言語化して記述する\n"
    "- テキストは原文のまま忠実に抽出する\n"
    "- 余計な説明や前置きは不要。抽出結果のみを返す"
)

# --- Excel prompt ---

EXCEL_FORMAT_PROMPT = (
    "以下はExcelスプレッドシートから抽出されたタブ区切りデータです。\n"
    "これをRAG（情報検索）に最適な構造化Markdownに整形してください。\n\n"
    "ルール:\n"
    "- テーブルデータはMarkdownテーブルに変換する\n"
    "- 見出し・セクション構造を適切に付与する\n"
    "- 注釈・備考は本文として記述する\n"
    "- 空の値やノイズ（重複した値はマージセルの痕跡）は除去する\n"
    "- データの意味や関係性を保持する\n"
    "- 余計な説明は不要。整形結果のみを返す"
)

_MEDIA_TYPES = {
    ".pdf": "application/pdf",
}

_FILE_TYPE_MAP = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".xlsx": "xlsx",
    ".xls": "xlsx",
    ".pptx": "pptx",
    ".ppt": "pptx",
    ".key": "key",
    ".txt": "txt",
    ".md": "md",
    ".csv": "csv",
    ".json": "json",
    ".html": "html",
    ".htm": "html",
}

# LibreOffice PDF変換で処理する拡張子
_LIBREOFFICE_EXTENSIONS = {".pptx", ".ppt", ".key"}

# PDF thresholds
TEXT_RICH_CHAR_THRESHOLD = 500
TEXT_RICH_IMAGE_LIMIT = 2
MAX_WORKERS = 5

# Excel thresholds
EXCEL_CHUNK_ROWS = 200
EXCEL_LARGE_FILE_MB = 50


class DocumentProcessor:
    def __init__(self):
        self._client: anthropic.Anthropic | None = None

    @property
    def client(self) -> anthropic.Anthropic:
        if self._client is None:
            self._client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
        return self._client

    # ================================================================
    # Common: Claude text formatting
    # ================================================================

    def _call_claude_text(self, text: str, prompt: str) -> str:
        message = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[{
                "role": "user",
                "content": f"{prompt}\n\n{text}",
            }],
        )
        return message.content[0].text

    # ================================================================
    # PDF extraction
    # ================================================================

    def _call_claude_document(self, file_content: bytes, media_type: str) -> str:
        encoded = base64.standard_b64encode(file_content).decode("utf-8")
        message = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8192,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": encoded,
                        },
                    },
                    {
                        "type": "text",
                        "text": DOCUMENT_EXTRACTION_PROMPT,
                    },
                ],
            }],
        )
        return message.content[0].text

    def _call_claude_vision(
        self, image_data: bytes, page_num: int, media_type: str = "image/png"
    ) -> str:
        encoded = base64.standard_b64encode(image_data).decode("utf-8")
        message = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": encoded,
                        },
                    },
                    {
                        "type": "text",
                        "text": PAGE_EXTRACTION_PROMPT,
                    },
                ],
            }],
        )
        return message.content[0].text

    def _classify_page(self, page: fitz.Page) -> str:
        text = page.get_text("text").strip()
        image_count = len(page.get_images(full=True))
        if len(text) >= TEXT_RICH_CHAR_THRESHOLD and image_count <= TEXT_RICH_IMAGE_LIMIT:
            return "text"
        return "vision"

    def _render_page_to_image(self, page: fitz.Page) -> tuple[bytes, str]:
        max_bytes = 3_700_000  # ~5MB after base64 encoding (×4/3)
        for dpi in (300, 200, 150):
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            data = pix.tobytes("png")
            if len(data) <= max_bytes:
                return data, "image/png"
            logger.debug("Page PNG at %d DPI = %d KB, reducing", dpi, len(data) // 1024)
        import io
        from PIL import Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        data = buf.getvalue()
        if len(data) <= max_bytes:
            return data, "image/jpeg"
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=60)
        return buf.getvalue(), "image/jpeg"

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        doc = fitz.open(stream=file_content, filetype="pdf")
        page_count = len(doc)
        logger.info("PDF extraction: %d pages, page-by-page hybrid mode", page_count)

        text_pages = []
        vision_pages = []
        for i in range(page_count):
            if self._classify_page(doc[i]) == "text":
                text_pages.append(i)
            else:
                vision_pages.append(i)

        logger.info(
            "Classification: %d text-rich (PyMuPDF), %d visual (Claude Vision)",
            len(text_pages), len(vision_pages),
        )

        results: dict[int, str] = {}

        for i in text_pages:
            results[i] = doc[i].get_text("text").strip()

        if vision_pages:
            page_images: list[tuple[int, bytes, str]] = []
            for i in vision_pages:
                img_data, media_type = self._render_page_to_image(doc[i])
                page_images.append((i, img_data, media_type))

            with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                futures = {
                    executor.submit(self._call_claude_vision, img_data, idx + 1, mt): idx
                    for idx, img_data, mt in page_images
                }
                for future in as_completed(futures):
                    idx = futures[future]
                    try:
                        text = future.result()
                        results[idx] = text
                        logger.info("Page %d: Vision extraction complete", idx + 1)
                    except Exception:
                        logger.exception("Page %d: Vision extraction failed", idx + 1)
                        results[idx] = doc[idx].get_text("text").strip()

        doc.close()

        parts = []
        for i in range(page_count):
            page_text = results.get(i, "")
            if page_text:
                parts.append(f"<!-- page {i + 1} -->\n{page_text}")

        return "\n\n".join(parts)

    # ================================================================
    # Excel extraction: openpyxl → Claude formatting
    # ================================================================

    @staticmethod
    def _sheet_to_tsv(ws) -> tuple[str, int]:
        """Extract worksheet to TSV text. Returns (tsv_text, row_count)."""
        lines: list[str] = []
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            cells = []
            for val in row:
                if val is None:
                    cells.append("")
                elif isinstance(val, float) and val == int(val):
                    cells.append(str(int(val)))
                else:
                    cells.append(str(val).replace("\t", " ").replace("\n", " "))
            if any(c.strip() for c in cells):
                lines.append("\t".join(cells))
                row_count += 1
        return "\n".join(lines), row_count

    def _format_excel_chunk(self, tsv_text: str, sheet_name: str, chunk_info: str = "") -> str:
        """Send a sheet chunk to Claude for formatting."""
        header = f"=== シート: {sheet_name} ==="
        if chunk_info:
            header += f" ({chunk_info})"
        return self._call_claude_text(f"{header}\n{tsv_text}", EXCEL_FORMAT_PROMPT)

    def extract_text_from_excel(self, file_content: bytes, filename: str) -> str:
        ext = Path(filename).suffix.lower()

        # Save to temp file for openpyxl
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name

        try:
            is_large = len(file_content) > EXCEL_LARGE_FILE_MB * 1024 * 1024

            if is_large:
                wb = openpyxl.load_workbook(tmp_path, read_only=True, data_only=True)
            else:
                wb = openpyxl.load_workbook(tmp_path, data_only=True)
                # Unmerge cells (only possible in non-read_only mode)
                for ws in wb.worksheets:
                    for cell_range in list(ws.merged_cells.ranges):
                        ws.unmerge_cells(str(cell_range))

            sheet_names = wb.sheetnames
            logger.info("Excel extraction: %s (%d sheets, %s)",
                        filename, len(sheet_names),
                        "read_only" if is_large else "standard")

            # Build table of contents
            toc_lines = ["このExcelファイルには以下のシートが含まれています:"]

            # Extract each sheet
            sheet_results: dict[int, str] = {}
            sheet_tasks: list[tuple[int, str, str, str]] = []  # (idx, name, tsv, chunk_info)

            for idx, name in enumerate(sheet_names):
                ws = wb[name]
                tsv_text, row_count = self._sheet_to_tsv(ws)
                toc_lines.append(f"- {name} ({row_count}行)")

                if not tsv_text.strip():
                    continue

                if row_count <= EXCEL_CHUNK_ROWS:
                    # Small sheet: single Claude call
                    sheet_tasks.append((idx, name, tsv_text, ""))
                else:
                    # Large sheet: split into chunks with header row
                    rows = tsv_text.split("\n")
                    header_row = rows[0]
                    data_rows = rows[1:]
                    chunk_num = 0
                    for start in range(0, len(data_rows), EXCEL_CHUNK_ROWS):
                        chunk_rows = data_rows[start:start + EXCEL_CHUNK_ROWS]
                        chunk_tsv = header_row + "\n" + "\n".join(chunk_rows)
                        chunk_num += 1
                        info = f"Part {chunk_num}, 行{start + 2}-{start + 1 + len(chunk_rows)}"
                        sheet_tasks.append((idx, name, chunk_tsv, info))

            wb.close()

            logger.info("Excel: %d sheets → %d Claude calls", len(sheet_names), len(sheet_tasks))

            # Process all sheet tasks in parallel
            with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                futures = {
                    executor.submit(
                        self._format_excel_chunk, tsv, name, info
                    ): (idx, name, info)
                    for idx, name, tsv, info in sheet_tasks
                }
                for future in as_completed(futures):
                    idx, name, info = futures[future]
                    key = (idx, info)
                    try:
                        result = future.result()
                        sheet_results[key] = result
                        logger.info("Sheet '%s' %s: formatted", name, info)
                    except Exception:
                        logger.exception("Sheet '%s' %s: formatting failed", name, info)
                        # Fallback: use raw TSV
                        for task_idx, task_name, task_tsv, task_info in sheet_tasks:
                            if task_idx == idx and task_info == info:
                                sheet_results[key] = task_tsv
                                break

            # Assemble in order
            parts = ["\n".join(toc_lines)]
            for idx, name, tsv, info in sheet_tasks:
                key = (idx, info)
                text = sheet_results.get(key, "")
                if text:
                    parts.append(f"<!-- sheet: {name} -->\n{text}")

            return "\n\n".join(parts)

        finally:
            Path(tmp_path).unlink(missing_ok=True)

    # ================================================================
    # Other formats
    # ================================================================

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        doc = docx.Document(io.BytesIO(file_content))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # 見出しスタイルをMarkdown化
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                prefix = "#" * (int(level) if level.isdigit() else 1)
                parts.append(f"{prefix} {text}")
            else:
                parts.append(text)
        # テーブル抽出
        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                # Markdownテーブル化
                header = "| " + " | ".join(rows[0]) + " |"
                separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
                body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
                parts.append(f"\n{header}\n{separator}\n{body}")
        return "\n\n".join(parts)

    def _convert_to_pdf_via_libreoffice(self, file_content: bytes, suffix: str) -> bytes:
        """LibreOfficeでPDFに変換し、既存PDFパイプラインで処理"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            src = Path(tmp_dir) / f"input{suffix}"
            src.write_bytes(file_content)
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmp_dir, str(src)],
                capture_output=True, timeout=120,
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr.decode()}")
            pdf_path = src.with_suffix(".pdf")
            if not pdf_path.exists():
                raise RuntimeError("LibreOffice did not produce a PDF")
            return pdf_path.read_bytes()

    def extract_text_via_pdf_conversion(self, file_content: bytes, suffix: str) -> str:
        """LibreOfficeでPDF変換 → 既存PDFパイプライン（PyMuPDF + Claude Vision）で抽出"""
        logger.info("Converting %s to PDF via LibreOffice for visual extraction", suffix)
        pdf_bytes = self._convert_to_pdf_via_libreoffice(file_content, suffix)
        return self.extract_text_from_pdf(pdf_bytes)

    @staticmethod
    def extract_text_from_csv(file_content: bytes) -> str:
        text = file_content.decode("utf-8-sig")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return ""
        # Markdownテーブル化
        header = "| " + " | ".join(rows[0]) + " |"
        separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
        body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:] if any(c.strip() for c in r))
        return f"{header}\n{separator}\n{body}"

    @staticmethod
    def extract_text_from_html(file_content: bytes) -> str:
        from bs4 import BeautifulSoup
        from markdownify import markdownify
        soup = BeautifulSoup(file_content, "html.parser")
        # bodyタグがあればそこだけ抽出
        body = soup.find("body")
        html_str = str(body) if body else str(soup)
        return markdownify(html_str, heading_style="ATX", strip=["script", "style"])

    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        # BOM付きUTF-8にも対応
        return file_content.decode("utf-8-sig")

    def extract_text(self, filename: str, file_content: bytes) -> str:
        ext = Path(filename).suffix.lower()
        # PPTX/KEY等 → LibreOfficeでPDF変換 → Claude Visionパイプライン
        if ext in _LIBREOFFICE_EXTENSIONS:
            return self.extract_text_via_pdf_conversion(file_content, ext)
        elif ext == ".pdf":
            return self.extract_text_from_pdf(file_content)
        elif ext in (".docx", ".doc"):
            return self.extract_text_from_docx(file_content)
        elif ext in (".xlsx", ".xls"):
            return self.extract_text_from_excel(file_content, filename)
        elif ext == ".csv":
            return self.extract_text_from_csv(file_content)
        elif ext in (".html", ".htm"):
            return self.extract_text_from_html(file_content)
        elif ext in (".txt", ".md", ".json"):
            return self.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    def get_file_type(filename: str) -> str:
        ext = Path(filename).suffix.lower()
        file_type = _FILE_TYPE_MAP.get(ext)
        if file_type is None:
            raise ValueError(f"Unsupported file type: {ext}")
        return file_type


document_processor = DocumentProcessor()
