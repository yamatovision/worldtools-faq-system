import json
import os
import re
from datetime import datetime, timedelta, timezone
from typing import AsyncGenerator

import anthropic
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from sqlalchemy import text, or_
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.document import ChatHistory, Document


TOOLS = [
    {
        "name": "get_quality_issues",
        "description": "回答失敗・低評価の質問を取得します。ナレッジのギャップ分析に使います。",
        "input_schema": {
            "type": "object",
            "properties": {
                "days": {
                    "type": "integer",
                    "description": "直近何日分を取得するか（デフォルト: 30）",
                    "default": 30,
                },
                "issue_type": {
                    "type": "string",
                    "description": "取得する問題の種類: 'no_answer'(回答失敗), 'bad'(低評価), 'all'(両方)",
                    "default": "all",
                },
            },
        },
    },
    {
        "name": "get_existing_documents",
        "description": "登録済みドキュメントの一覧を取得します。ファイル名・カテゴリ・チャンク数を確認できます。",
        "input_schema": {
            "type": "object",
            "properties": {},
        },
    },
    {
        "name": "get_document_content",
        "description": "指定ドキュメントの全文を取得します。既存ドキュメントの内容確認に使います。",
        "input_schema": {
            "type": "object",
            "properties": {
                "document_id": {
                    "type": "string",
                    "description": "ドキュメントID",
                },
            },
            "required": ["document_id"],
        },
    },
    {
        "name": "generate_document",
        "description": "Markdown形式のコンテンツからWord文書(.docx)を生成し、ダウンロードリンクを返します。管理者の明確な指示があってから使ってください。",
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "ドキュメントのタイトル",
                },
                "content_markdown": {
                    "type": "string",
                    "description": "ドキュメントの本文（Markdown形式）",
                },
            },
            "required": ["title", "content_markdown"],
        },
    },
]

SYSTEM_PROMPT = """あなたは社内ナレッジベースの品質改善を支援する管理者向けアシスタントです。

【あなたの役割】
- 回答品質データ（回答失敗・低評価）を分析して、ナレッジベースのギャップを特定する
- 不足しているドキュメントや補足すべき内容を提案する
- 管理者との対話を通じて、新しいドキュメントの内容を詰める
- 管理者の指示があれば、Word文書を生成する

【行動手順】
1. まず get_quality_issues で品質データを取得し、傾向を分析する
2. get_existing_documents で既存のナレッジ範囲を把握する
3. ギャップを特定し、具体的なドキュメント作成提案を行う
4. 管理者と内容を詰めた後、generate_document でWord文書を生成する

【重要なルール】
- generate_document は管理者の明確な指示（「作成して」「生成して」等）があってから使う
- 分析結果は簡潔にまとめ、箇条書きで見やすくする
- 既存ドキュメントとの重複がないか確認する
- 日本語で回答する"""


def _sse(data: dict) -> str:
    return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"


def _content_to_dict(content_blocks) -> list[dict]:
    """Anthropic SDK ContentBlock objects to serializable dicts."""
    result = []
    for block in content_blocks:
        if block.type == "text":
            result.append({"type": "text", "text": block.text})
        elif block.type == "tool_use":
            result.append({
                "type": "tool_use",
                "id": block.id,
                "name": block.name,
                "input": block.input,
            })
    return result


def _markdown_to_docx(title: str, markdown: str) -> str:
    """MarkdownテキストからWord文書を生成し、ファイルパスを返す。"""
    doc = DocxDocument()

    # タイトル
    title_para = doc.add_heading(title, level=0)
    title_para.runs[0].font.size = Pt(18)

    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 見出し
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        # 箇条書き
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        # 番号付きリスト
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line).strip()
            doc.add_paragraph(text, style="List Number")
        # 空行
        elif not line.strip():
            pass
        # 通常テキスト
        else:
            para = doc.add_paragraph()
            # 太字処理
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)

        i += 1

    # 保存
    output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "generated_docs")
    os.makedirs(output_dir, exist_ok=True)

    safe_title = re.sub(r'[^\w\s\u3000-\u9fff-]', '', title)[:50].strip()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_title}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)

    doc.save(filepath)
    return filename


TOOL_LABELS = {
    "get_quality_issues": "品質データ取得中...",
    "get_existing_documents": "ドキュメント一覧取得中...",
    "get_document_content": "ドキュメント内容確認中...",
    "generate_document": "Word文書生成中...",
}


class AdminAgent:
    def __init__(self, db: Session, organization_id: str):
        self.db = db
        self.organization_id = organization_id
        self.client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

    def _execute_tool(self, name: str, input_data: dict) -> str:
        if name == "get_quality_issues":
            return self._tool_get_quality_issues(
                input_data.get("days", 30),
                input_data.get("issue_type", "all"),
            )
        elif name == "get_existing_documents":
            return self._tool_get_existing_documents()
        elif name == "get_document_content":
            return self._tool_get_document_content(input_data["document_id"])
        elif name == "generate_document":
            return self._tool_generate_document(
                input_data["title"],
                input_data["content_markdown"],
            )
        return json.dumps({"error": f"Unknown tool: {name}"})

    def _tool_get_quality_issues(self, days: int = 30, issue_type: str = "all") -> str:
        since = datetime.now(timezone.utc) - timedelta(days=days)
        query = self.db.query(ChatHistory).filter(
            ChatHistory.organization_id == self.organization_id,
            ChatHistory.created_at >= since,
        )

        if issue_type == "no_answer":
            query = query.filter(ChatHistory.is_no_answer == "1")
        elif issue_type == "bad":
            query = query.filter(ChatHistory.feedback == "bad")
        else:
            query = query.filter(
                or_(ChatHistory.is_no_answer == "1", ChatHistory.feedback == "bad")
            )

        results = query.order_by(ChatHistory.created_at.desc()).limit(50).all()
        items = []
        for r in results:
            items.append({
                "question": r.question,
                "answer": r.answer[:200] if r.answer else "",
                "is_no_answer": r.is_no_answer == "1",
                "feedback": r.feedback,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            })

        return json.dumps({
            "period_days": days,
            "total": len(items),
            "items": items,
        }, ensure_ascii=False)

    def _tool_get_existing_documents(self) -> str:
        sql = text("""
            SELECT d.id, d.filename, d.category,
                   COUNT(dc.id) as chunk_count
            FROM documents d
            LEFT JOIN document_chunks dc ON d.id = dc.document_id
            WHERE d.organization_id = :org_id
            GROUP BY d.id, d.filename, d.category
            ORDER BY d.filename
        """)
        rows = self.db.execute(sql, {"org_id": self.organization_id}).fetchall()
        docs = [
            {
                "id": row.id,
                "filename": row.filename,
                "category": row.category or "",
                "chunk_count": row.chunk_count,
            }
            for row in rows
        ]
        return json.dumps({"documents": docs, "total": len(docs)}, ensure_ascii=False)

    def _tool_get_document_content(self, document_id: str) -> str:
        sql = text("""
            SELECT dc.content, dc.chunk_index, d.filename
            FROM document_chunks dc
            JOIN documents d ON dc.document_id = d.id
            WHERE dc.document_id = :document_id
              AND dc.organization_id = :organization_id
            ORDER BY dc.chunk_index
        """)
        rows = self.db.execute(sql, {
            "document_id": document_id,
            "organization_id": self.organization_id,
        }).fetchall()
        if not rows:
            return json.dumps({"error": "ドキュメントが見つかりません。"}, ensure_ascii=False)
        filename = rows[0].filename
        full_text = "\n\n".join(row.content for row in rows)
        return json.dumps({
            "filename": filename,
            "content": full_text,
            "chunk_count": len(rows),
        }, ensure_ascii=False)

    def _tool_generate_document(self, title: str, content_markdown: str) -> str:
        filename = _markdown_to_docx(title, content_markdown)
        return json.dumps({
            "status": "ok",
            "filename": filename,
            "message": f"Word文書「{filename}」を生成しました。",
        }, ensure_ascii=False)

    async def run(self, message: str, conversation_history: list[dict]) -> AsyncGenerator[str, None]:
        messages = []
        for msg in conversation_history[-20:]:
            messages.append({"role": msg["role"], "content": msg["content"]})
        messages.append({"role": "user", "content": message})

        system_with_cache = [
            {"type": "text", "text": SYSTEM_PROMPT, "cache_control": {"type": "ephemeral"}},
        ]
        tools_with_cache = TOOLS[:-1] + [{**TOOLS[-1], "cache_control": {"type": "ephemeral"}}]

        max_iterations = 15
        for i in range(max_iterations):
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4096,
                temperature=0.3,
                system=system_with_cache,
                tools=tools_with_cache,
                messages=messages,
            )

            if response.stop_reason != "tool_use":
                for block in response.content:
                    if block.type == "text":
                        text_content = block.text
                        chunk_size = 4
                        for pos in range(0, len(text_content), chunk_size):
                            yield _sse({"token": text_content[pos:pos + chunk_size]})
                break

            assistant_content = response.content
            tool_results = []

            for block in assistant_content:
                if block.type != "tool_use":
                    continue

                tool_name = block.name
                tool_input = block.input

                yield _sse({"step": {
                    "tool": tool_name,
                    "status": "running",
                    "label": TOOL_LABELS.get(tool_name, f"{tool_name} 実行中..."),
                }})

                result = self._execute_tool(tool_name, tool_input)
                summary = self._summarize_result(tool_name, result)

                step_event: dict = {
                    "tool": tool_name,
                    "status": "done",
                    "summary": summary,
                }

                # generate_document の場合、ダウンロード情報を付与
                if tool_name == "generate_document":
                    try:
                        result_data = json.loads(result)
                        if result_data.get("status") == "ok":
                            yield _sse({"download": {
                                "filename": result_data["filename"],
                            }})
                    except json.JSONDecodeError:
                        pass

                yield _sse({"step": step_event})

                tool_results.append({
                    "type": "tool_result",
                    "tool_use_id": block.id,
                    "content": result,
                })

            messages.append({"role": "assistant", "content": _content_to_dict(assistant_content)})
            messages.append({"role": "user", "content": tool_results})
        else:
            yield _sse({"token": "処理が長くなっています。ここまでの情報をもとに回答します。"})

        yield _sse({"done": True})

    @staticmethod
    def _summarize_result(tool_name: str, result_json: str) -> str:
        try:
            data = json.loads(result_json)
        except json.JSONDecodeError:
            return "結果を取得"

        if tool_name == "get_quality_issues":
            return f"{data.get('total', 0)}件の品質問題を取得"
        elif tool_name == "get_existing_documents":
            return f"{data.get('total', 0)}件のドキュメント一覧を取得"
        elif tool_name == "get_document_content":
            if "error" in data:
                return "ドキュメントが見つかりません"
            return f"「{data.get('filename', '')}」の全文を取得"
        elif tool_name == "generate_document":
            if data.get("status") == "ok":
                return f"Word文書「{data.get('filename', '')}」を生成"
            return "生成に失敗"
        return "結果を取得"
